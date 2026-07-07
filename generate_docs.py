#!/usr/bin/env python3
"""Generate BopheloHub documentation.docx"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

S=RGBColor; H=WD_ALIGN_PARAGRAPH.CENTER; A=WD_TABLE_ALIGNMENT.CENTER

def sc(c,h):
    e=OxmlElement('w:shd'); e.set(qn('w:fill'),h); c._tc.get_or_add_tcPr().append(e)

def cd(d,t):
    p=d.add_paragraph(); p.style=d.styles['No Spacing']
    r=p.add_run(t); r.font.name='Courier New'; r.font.size=Pt(8)
    r.font.color.rgb=S(0x1A,0x1A,0x2E)
    s=OxmlElement('w:shd'); s.set(qn('w:fill'),'F0F0F0')
    p._element.get_or_add_pPr().append(s)

def bu(d,t,l=0,p=None):
    p2=d.add_paragraph(style='List Bullet')
    p2.paragraph_format.left_indent=Cm(1.27+l*1.27)
    if p: r=p2.add_run(p); r.bold=True; p2.add_run(t)
    else: p2.add_run(t)

def tr(t,d,h=False):
    r=t.add_row()
    for i,x in enumerate(d):
        c=r.cells[i]; c.text=str(x)
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT
            for r2 in p.runs: r2.font.size=Pt(9)
            if h: r2.bold=True
        if h: sc(c,'2B579A')
        if h:
            for p in c.paragraphs:
                for r2 in p.runs: r2.font.color.rgb=S(0xFF,0xFF,0xFF)

d=Document(); d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(11)

# Cover
for _ in range(6): d.add_paragraph('')
t=d.add_paragraph(); t.alignment=H
r=t.add_run('BopheloHub'); r.font.size=Pt(42); r.bold=True; r.font.color.rgb=S(0x1A,0x1A,0x2E)
t=d.add_paragraph(); t.alignment=H
r=t.add_run('E-Learning Platform'); r.font.size=Pt(24); r.font.color.rgb=S(0xE9,0x45,0x60)
t=d.add_paragraph(); t.alignment=H
r=t.add_run('Empowering Lives Through Education'); r.font.size=Pt(16); r.font.color.rgb=S(0x66,0x66,0x66)
d.add_paragraph('')
t=d.add_paragraph(); t.alignment=H
r=t.add_run('Project Documentation\nGenerated: June 2026'); r.font.size=Pt(12); r.font.color.rgb=S(0x99,0x99,0x99)
# TOC
d.add_heading('Table of Contents',level=1)
for i in ['1. Project Overview','2. Technology Stack','3. Project Structure',
    '4. Configuration','5. Authentication','6. Courses','7. Lessons',
    '8. Enrollments','9. Payments','10. Certificates','11. Quizzes',
    '12. Reviews','13. Discussions','14. Notifications','15. Instructors',
    '16. Owner','17. Analytics','18. URLs','19. Security','20. Deployment']:
    p=d.add_paragraph(i); p.paragraph_format.space_after=Pt(2); p.runs[0].font.size=Pt(11)
d.add_page_break()

# 1
d.add_heading('1. Project Overview',level=1)
d.add_paragraph('BopheloHub is a Django (6.0.5) e-learning platform connecting students, instructors, and administrators. Features: course creation, CKEditor 5 lessons, video embedding, free/paid enrollment, payment verification, quizzes, assessments, PDF certificates (ReportLab), discussion forums, AJAX reviews, notifications, and dashboards.')
d.add_heading('Key Features',level=2)
for p,t in [('User Roles: ','Students, Instructors, Platform Owner'),
    ('Courses: ','Create/edit/publish with categories & thumbnails'),
    ('Curriculum: ','Sections with drag-drop reordering'),
    ('Lessons: ','CKEditor 5, video, resources, notes, comments'),
    ('Enrollments: ','Free & paid with payment proof upload'),
    ('Payments: ','Manual verify, 85/15 commission split'),
    ('Quizzes: ','Per-lesson MCQ with auto-scoring, 70% pass'),
    ('Assessments: ','Course-level MCQ & essay, manual grading'),
    ('Certificates: ','Auto PDF via ReportLab, UUID codes'),
    ('Forums: ','Course discussions, pinned threads, replies'),
    ('Reviews: ','5-star AJAX rating, enrolled-only'),
    ('Notifications: ','In-app alerts for events'),
    ('Dashboards: ','Instructor & Owner analytics')]:
    bu(d,t,p=p)
d.add_page_break()

# 2
d.add_heading('2. Technology Stack',level=1)
tb=d.add_table(rows=1,cols=3); tb.style='Table Grid'; tb.alignment=A
tr(tb,['Layer','Technology','Version'],h=True)
for r in [['Backend','Django','6.0.5'],['Language','Python','3.13'],
    ['Database','PostgreSQL/Supabase','psycopg2-binary'],
    ['Server','Gunicorn+WhiteNoise','26.0.0'],
    ['Rich Text','django-ckeditor-5','0.2.20'],
    ['Storage','Cloudinary','latest'],['PDF','ReportLab','4.5.1'],
    ['Forms','crispy-bootstrap5','2026.3'],
    ['Frontend','Bootstrap 5','CDN'],['Config','python-decouple','3.8']]:
    tr(tb,r)
d.add_paragraph('')
d.add_heading('Key Dependencies',level=2)
for dep in ['Django==6.0.5','django-ckeditor-5==0.2.20',
    'django-crispy-forms==2.6','crispy-bootstrap5==2026.3',
    'cloudinary>=1.4.0','django-cloudinary-storage==0.3.0',
    'gunicorn==26.0.0','psycopg2-binary==2.9.12',
    'python-decouple==3.8','reportlab==4.5.1',
    'whitenoise==6.12.0','pillow==12.2.0']:
    bu(d,dep)
# 3
d.add_heading('3. Project Structure',level=1)
cd(d,'''elearning/
  manage.py              # Django CLI
  config/                # Settings, URLs, ASGI/WSGI
  apps/
    accounts/            # User auth & profiles
    instructors/         # Verification
    courses/             # Course mgmt
    lessons/             # Sections, lessons, resources
    enrollments/         # Enrollment & progress
    payments/            # Payment verification
    certificates/        # PDF generation
    quizzes/             # Quiz engine
    reviews/             # Course ratings
    assessments/         # Course assessments
    discussions/         # Forums
    notifications/       # In-app alerts
    analytics/           # View tracking
    owner/               # Admin dashboard
  templates/             # HTML templates
  static/                # CSS, images
  media/                 # Uploaded files''')
d.add_page_break()

# 4
d.add_heading('4. Configuration & Environment',level=1)
d.add_paragraph('Environment via .env file (python-decouple):')
tl=d.add_table(rows=1,cols=3); tl.style='Table Grid'; tl.alignment=A
tr(tl,['Variable','Description','Example'],h=True)
for r in [['SECRET_KEY','Django secret','insecure-...'],
    ['DEBUG','Debug mode','True/False'],
    ['ALLOWED_HOSTS','Allowed hosts','localhost,127.0.0.1'],
    ['DB_NAME/DB_USER','DB name/user','postgres/...'],
    ['DB_HOST','Database host','supabase.com'],
    ['CLOUDINARY_*','Cloudinary config','key + secret'],
    ['EMAIL_HOST','SMTP','smtp.gmail.com']]:
    tr(tl,r)
d.add_page_break()

# 5
d.add_heading('5. Authentication (accounts)',level=1)
d.add_heading('Model: User',level=2)
d.add_paragraph('Extends AbstractUser. Custom: email (unique, USERNAME_FIELD), is_instructor, bio, profile_picture. Username auto-generated from email.')
d.add_heading('Views',level=2)
tl=d.add_table(rows=1,cols=3); tl.style='Table Grid'; tl.alignment=A
tr(tl,['View','URL','Description'],h=True)
for r in [['StudentRegisterView','/register/student/','Reg + auto-login'],
    ['InstructorRegisterView','/register/instructor/','Reg + instructor apply'],
    ['CustomLoginView','/login/','Role-based redirect'],
    ['ProfileView','/profile/','Edit profile, view enrollments/certs']]:
    tr(tl,r)
d.add_paragraph('')
d.add_heading('Forms',level=2)
bu(d,'StudentRegistrationForm: email, first_name, last_name, password1, password2')
bu(d,'InstructorRegistrationForm: same + is_instructor=True')
bu(d,'ProfileEditForm: first_name, last_name, bio, profile_picture')
d.add_page_break()


